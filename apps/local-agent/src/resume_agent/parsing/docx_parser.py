"""DOCX paragraph and table text extraction."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from .models import ParsedSegment


def parse_docx(path: str | Path) -> list[ParsedSegment]:
    document = Document(str(path))
    segments: list[ParsedSegment] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            segments.append(ParsedSegment(text=text, location=f"paragraph {index}"))
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                text = cell.text.strip()
                if text:
                    segments.append(
                        ParsedSegment(
                            text=text,
                            location=(f"table {table_index} row {row_index} cell {cell_index}"),
                            extraction_method="table",
                        )
                    )
    return segments
