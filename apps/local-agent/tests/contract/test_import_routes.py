import base64
import json
from pathlib import Path

from fastapi.testclient import TestClient
from jsonschema import Draft202012Validator, FormatChecker

from resume_agent.api.app import create_app
from resume_agent.profile.service import ProfileService

ROOT = Path(__file__).resolve().parents[4]


def _validator(name: str) -> Draft202012Validator:
    master = json.loads(
        (ROOT / "packages/contracts/v0.1/contracts.schema.json").read_text(encoding="utf-8")
    )
    root = {
        "$schema": master["$schema"],
        "$id": master["$id"],
        "$defs": master["$defs"],
        "$ref": f"#/$defs/{name}",
    }
    return Draft202012Validator(root, format_checker=FormatChecker())


def test_import_preview_uses_local_content_and_does_not_write_profile(fake_profile_store) -> None:
    app = create_app(profile_service=ProfileService(fake_profile_store), require_loopback=False)
    client = TestClient(app)
    payload = {
        "schema_version": "0.1",
        "request_id": "req-import-route-1",
        "task_id": "task-import-route-1",
        "operation": "profile.import.preview",
        "source": {
            "document_id": "doc-import-1",
            "filename": "resume.docx",
            "media_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        },
        "content_base64": base64.b64encode(b"not a docx").decode(),
        "consent": {"remote_model_allowed": False},
        "ocr_mode": "auto",
    }

    response = client.post("/v0/profile/import/preview", json=payload)

    assert response.status_code == 400
    assert response.json()["error"]["code"] in {"DOCUMENT_PARSE_FAILED", "INVALID_FIELD_VALUE"}
    assert fake_profile_store.write_calls == 0


def test_import_preview_and_confirm_round_trip(fake_profile_store, tmp_path) -> None:
    from io import BytesIO

    from docx import Document

    document = Document()
    document.add_paragraph("姓名：示例用户 邮箱 example@example.test")
    buffer = BytesIO()
    document.save(buffer)
    content = buffer.getvalue()
    app = create_app(profile_service=ProfileService(fake_profile_store), require_loopback=False)
    client = TestClient(app)
    preview = client.post(
        "/v0/profile/import/preview",
        json={
            "schema_version": "0.1",
            "request_id": "req-roundtrip",
            "task_id": "task-roundtrip",
            "operation": "profile.import.preview",
            "source": {
                "document_id": "doc-roundtrip",
                "filename": "resume.docx",
                "media_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "size_bytes": len(content),
            },
            "content_base64": base64.b64encode(content).decode(),
            "consent": {"remote_model_allowed": False},
        },
    )
    assert preview.status_code == 200
    _validator("ProfileImportPreview").validate(preview.json())
    candidate = next(
        item for item in preview.json()["candidates"] if item["field_id"] == "contact.email"
    )

    confirmed = client.post(
        "/v0/profile/import/confirm",
        json={
            "schema_version": "0.1",
            "request_id": "req-confirm",
            "task_id": "task-roundtrip",
            "operation": "profile.import.confirm",
            "profile_id": "default-profile",
            "expected_profile_version": 0,
            "decisions": [
                {
                    "candidate_id": candidate["candidate_id"],
                    "decision": "accept",
                    "user_confirmed": True,
                }
            ],
        },
    )
    assert confirmed.status_code == 200
    _validator("ProfileImportConfirmResponse").validate(confirmed.json())
    assert confirmed.json()["written_field_ids"] == ["contact.email"]
    assert fake_profile_store.snapshot.fields[0].value == "example@example.test"
    replay = client.post(
        "/v0/profile/import/confirm",
        json={
            "schema_version": "0.1",
            "request_id": "req-confirm",
            "task_id": "task-roundtrip",
            "operation": "profile.import.confirm",
            "profile_id": "default-profile",
            "expected_profile_version": 0,
            "decisions": [
                {
                    "candidate_id": candidate["candidate_id"],
                    "decision": "accept",
                    "user_confirmed": True,
                }
            ],
        },
    )
    assert replay.status_code == 200
    assert replay.json() == confirmed.json()


def test_import_cancel_invalidates_preview_task(fake_profile_store) -> None:
    from io import BytesIO

    from docx import Document

    document = Document()
    document.add_paragraph("邮箱 example@example.test")
    buffer = BytesIO()
    document.save(buffer)
    app = create_app(profile_service=ProfileService(fake_profile_store), require_loopback=False)
    client = TestClient(app)
    preview = client.post(
        "/v0/profile/import/preview",
        json={
            "schema_version": "0.1",
            "request_id": "req-cancel-preview",
            "task_id": "task-cancel-preview",
            "operation": "profile.import.preview",
            "source": {
                "document_id": "doc-cancel",
                "filename": "resume.docx",
                "media_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            },
            "content_base64": base64.b64encode(buffer.getvalue()).decode(),
            "consent": {"remote_model_allowed": False},
        },
    )
    candidate = preview.json()["candidates"][0]
    cancelled = client.post(
        "/v0/profile/import/cancel",
        json={
            "schema_version": "0.1",
            "request_id": "req-cancel",
            "task_id": "task-cancel-preview",
            "operation": "profile.import.cancel",
        },
    )
    assert cancelled.status_code == 200
    confirm = client.post(
        "/v0/profile/import/confirm",
        json={
            "schema_version": "0.1",
            "request_id": "req-after-cancel",
            "task_id": "task-cancel-preview",
            "operation": "profile.import.confirm",
            "profile_id": "default-profile",
            "expected_profile_version": 0,
            "decisions": [
                {
                    "candidate_id": candidate["candidate_id"],
                    "decision": "accept",
                    "user_confirmed": True,
                }
            ],
        },
    )
    assert confirm.status_code == 400
    assert fake_profile_store.write_calls == 0
