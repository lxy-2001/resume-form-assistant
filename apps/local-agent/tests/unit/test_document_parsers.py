from pathlib import Path

from docx import Document
from pypdf import PdfWriter

from resume_agent.parsing.docx_parser import parse_docx
from resume_agent.parsing.pdf_parser import parse_pdf


def test_parse_pdf_returns_page_located_segments(tmp_path: Path) -> None:
    path = tmp_path / "resume.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=300, height=300)
    with path.open("wb") as stream:
        writer.write(stream)

    segments = parse_pdf(path)

    assert segments == []


def test_parse_docx_extracts_paragraphs_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("姓名：示例用户")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "邮箱"
    table.cell(0, 1).text = "example@example.test"
    document.save(path)

    segments = parse_docx(path)

    assert [segment.text for segment in segments] == [
        "姓名：示例用户",
        "邮箱",
        "example@example.test",
    ]
    assert segments[0].location == "paragraph 1"
    assert segments[1].location == "table 1 row 1 cell 1"
